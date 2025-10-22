"""
Generador de calendarios para Green Logistics
Toma datos de clientes de la base de datos y genera calendarios en Word
"""

import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import calendar
import os
import tempfile
import zipfile
from database import get_clients, get_calculated_dates, get_client_activities, get_frequency_template_by_id

# Mapeo de meses en español (independiente del locale del sistema)
MESES_ES = {
    1: 'ene', 2: 'feb', 3: 'mar', 4: 'abr', 5: 'may', 6: 'jun',
    7: 'jul', 8: 'ago', 9: 'sep', 10: 'oct', 11: 'nov', 12: 'dic'
}

def format_date_spanish(date_obj):
    """
    Formatea una fecha en formato español: 01-ene, 15-mar, etc.
    
    Args:
        date_obj: datetime object
    
    Returns:
        str: Fecha formateada en español (ej: '01-ene')
    """
    if not isinstance(date_obj, datetime):
        return ''
    
    day = date_obj.day
    month = MESES_ES.get(date_obj.month, 'xxx')
    return f"{day:02d}-{month}"

class CalendarGenerator:
    def __init__(self, template_path=None):
        self.template_path = template_path or "CF PLANTILLA CALENDARIO 2025.docx"
        
    def generate_dates_for_client(self, client_data, year=2025):
        """
        Genera fechas de envío OC y entrega basadas en la frecuencia del cliente
        """
        # Intentar obtener frecuencia del cliente o usar mensual por defecto
        frequency = client_data.get('frecuencia') or client_data.get('frequency')
        
        # Si no hay frecuencia, inferir del calendario_sap o usar mensual
        if not frequency:
            calendario_sap = client_data.get('calendario_sap', '')
            if calendario_sap:
                # Inferir frecuencia basada en el código SAP
                if calendario_sap.startswith('Q') or '15' in calendario_sap:
                    frequency = 'Quincenal'
                elif calendario_sap.startswith('M') or 'M' in calendario_sap:
                    frequency = 'Mensual'
                elif calendario_sap.startswith('B') or '60' in calendario_sap:
                    frequency = 'Bimensual'
                elif calendario_sap.startswith('T') or '90' in calendario_sap:
                    frequency = 'Trimestral'
                elif calendario_sap.startswith('S') or '180' in calendario_sap:
                    frequency = 'Semestral'
                elif calendario_sap.startswith('A') or '365' in calendario_sap:
                    frequency = 'Anual'
                else:
                    frequency = 'Mensual'  # Por defecto
            else:
                frequency = 'Mensual'  # Por defecto si no hay información
        
        start_date = datetime(year, 1, 1)
        
        dates_data = []
        
        if frequency == 'Quincenal':
            # Cada 15 días
            interval_days = 15
            lead_time = 6  # 6 días entre envío OC y entrega
        elif frequency == 'Mensual':
            # Cada 30 días
            interval_days = 30
            lead_time = 6
        elif frequency == 'Bimensual':
            # Cada 60 días
            interval_days = 60
            lead_time = 6
        elif frequency == 'Trimestral':
            # Cada 90 días
            interval_days = 90
            lead_time = 6
        elif frequency == 'Semestral':
            # Cada 180 días
            interval_days = 180
            lead_time = 10
        elif frequency == 'Anual':
            # Una vez al año
            interval_days = 365
            lead_time = 15
        else:
            # Default mensual
            interval_days = 30
            lead_time = 6
        
        current_date = start_date
        end_date = datetime(year, 12, 31)
        
        while current_date <= end_date:
            # Fecha de envío OC
            oc_date = current_date
            # Fecha de entrega (siempre después de OC)
            entrega_date = current_date + timedelta(days=lead_time)
            
            # Asegurar que no pasemos del año
            if entrega_date <= end_date:
                dates_data.append({
                    'Fecha envío OC': format_date_spanish(oc_date),
                    'Fecha Entrega': format_date_spanish(entrega_date)
                })
            
            # Siguiente fecha
            current_date += timedelta(days=interval_days)
        
        return dates_data
    
    def get_client_real_frequencies(self, client_id):
        """
        Obtiene las frecuencias reales del cliente desde la base de datos
        """
        try:
            activities_df = get_client_activities(client_id)
            if activities_df.empty:
                return {}
            
            frequencies = {}
            for _, activity in activities_df.iterrows():
                activity_name = activity['activity_name']
                frequency_template_id = activity['frequency_template_id']
                
                # Obtener detalles de la frecuencia
                frequency_template = get_frequency_template_by_id(frequency_template_id)
                if frequency_template is not None:
                    frequencies[activity_name] = {
                        'template_id': frequency_template_id,
                        'name': frequency_template['name'] if 'name' in frequency_template else 'Mensual',
                        'type': frequency_template['frequency_type'] if 'frequency_type' in frequency_template else 'specific_days',
                        'config': frequency_template['frequency_config'] if 'frequency_config' in frequency_template else '{}'
                    }
            
            return frequencies
            
        except Exception as e:
            print(f"Error obteniendo frecuencias reales del cliente {client_id}: {e}")
            return {}
    
    def generate_dates_with_real_frequencies(self, client_data, year=2026):
        """
        Genera fechas usando las frecuencias reales del cliente desde la BD
        """
        client_id = client_data.get('id')
        if not client_id:
            return self.generate_dates_for_client(client_data, year)
        
        try:
            # Obtener frecuencias reales del cliente
            real_frequencies = self.get_client_real_frequencies(client_id)
            
            if not real_frequencies:
                # Si no hay frecuencias en BD, usar método algorítmico
                return self.generate_dates_for_client(client_data, year)
            
            # Generar fechas basadas en las frecuencias reales
            all_dates = []
            
            # Buscar frecuencia de "Fecha Envío OC" para las fechas principales
            oc_frequency_info = None
            if 'Fecha Envío OC' in real_frequencies:
                oc_frequency_info = real_frequencies['Fecha Envío OC']
            elif real_frequencies:
                # Tomar la primera frecuencia disponible
                oc_frequency_info = list(real_frequencies.values())[0]
            
            if not oc_frequency_info:
                return self.generate_dates_for_client(client_data, year)
            
            # Interpretar la frecuencia para generar fechas específicas
            frequency_name = oc_frequency_info['name']
            
            # Generar fechas basándose en el patrón específico
            oc_dates = self.calculate_specific_dates_from_frequency(frequency_name, year)
            
            # Crear fechas de entrega (6 días después de OC)
            all_dates = []
            for oc_date in oc_dates:
                entrega_date = oc_date + timedelta(days=6)
                all_dates.append({
                    'Fecha envío OC': format_date_spanish(oc_date),
                    'Fecha Entrega': format_date_spanish(entrega_date)
                })
            
            return all_dates
            
        except Exception as e:
            print(f"Error generando fechas con frecuencias reales: {e}")
            return self.generate_dates_for_client(client_data, year)
    
    def calculate_specific_dates_from_frequency(self, frequency_name, year):
        """
        Calcula fechas específicas basándose en el patrón de frecuencia
        """
        dates = []
        
        # Mapear días de la semana
        weekdays = {
            'Lunes': 0, 'Martes': 1, 'Miércoles': 2, 
            'Jueves': 3, 'Viernes': 4, 'Sábado': 5, 'Domingo': 6
        }
        
        # Extraer información del patrón
        target_weekday = None
        weeks_in_month = []
        
        # Identificar el día de la semana
        for day_name, day_num in weekdays.items():
            if day_name in frequency_name:
                target_weekday = day_num
                break
        
        if target_weekday is None:
            # Si no encuentra patrón específico, usar intervalos simples
            return self.generate_simple_interval_dates(frequency_name, year)
        
        # Identificar qué semanas del mes
        if '1er y 3er' in frequency_name:
            weeks_in_month = [1, 3]
        elif '2do y 4to' in frequency_name:
            weeks_in_month = [2, 4]
        elif '1er 2do y 3er' in frequency_name:
            weeks_in_month = [1, 2, 3]
        elif '1er' in frequency_name:
            weeks_in_month = [1]
        elif '2do' in frequency_name:
            weeks_in_month = [2]
        elif '3er' in frequency_name:
            weeks_in_month = [3]
        elif '4to' in frequency_name:
            weeks_in_month = [4]
        else:
            # Por defecto, primera semana del mes
            weeks_in_month = [1]
        
        # Generar fechas para cada mes del año
        for month in range(1, 13):
            for week_num in weeks_in_month:
                # Encontrar el día específico de la semana en la semana específica del mes
                specific_date = self.find_nth_weekday_of_month(year, month, target_weekday, week_num)
                if specific_date:
                    dates.append(specific_date)
        
        return sorted(dates)
    
    def find_nth_weekday_of_month(self, year, month, weekday, n):
        """
        Encuentra el n-ésimo día de la semana específico en un mes
        weekday: 0=Lunes, 1=Martes, etc.
        n: 1=primero, 2=segundo, etc.
        """
        # Primer día del mes
        first_day = datetime(year, month, 1)
        
        # Encontrar el primer día de la semana objetivo en el mes
        days_ahead = weekday - first_day.weekday()
        if days_ahead < 0:  # El día objetivo ya pasó esta semana
            days_ahead += 7
        
        # Calcular la fecha del primer día objetivo
        first_occurrence = first_day + timedelta(days=days_ahead)
        
        # Calcular la fecha del n-ésimo día objetivo
        target_date = first_occurrence + timedelta(weeks=n-1)
        
        # Verificar que la fecha esté dentro del mes
        if target_date.month == month:
            return target_date
        else:
            return None
    
    def generate_simple_interval_dates(self, frequency_name, year):
        """
        Genera fechas con intervalos simples cuando no se puede determinar patrón específico
        """
        dates = []
        
        # Determinar intervalo basado en frecuencia
        if 'cada semana' in frequency_name.lower():
            interval_days = 7
        elif '1er y 3er' in frequency_name or '2do y 4to' in frequency_name:
            interval_days = 14
        else:
            interval_days = 30
        
        # Generar fechas con el intervalo
        start_date = datetime(year, 1, 1)
        end_date = datetime(year, 12, 31)
        current_date = start_date
        
        while current_date <= end_date:
            dates.append(current_date)
            current_date += timedelta(days=interval_days)
        
        return dates
    
    def generate_dates_from_database(self, client_data, year=2026):
        """
        Genera fechas ÚNICAMENTE desde la base de datos
        NO genera fechas algorítmicamente - solo lee lo que está en BD
        """
        client_id = client_data.get('id')
        if not client_id:
            print(f"⚠️  Cliente sin ID - no se pueden obtener fechas de la BDD")
            return []
        
        try:
            # Obtener fechas calculadas desde la base de datos
            calculated_dates_df = get_calculated_dates(client_id)
            
            if calculated_dates_df.empty:
                print(f"⚠️  No hay fechas en la BD para el cliente {client_id}")
                return []
            
            # Filtrar fechas para el año específico
            calculated_dates_df['date'] = pd.to_datetime(calculated_dates_df['date'])
            year_dates = calculated_dates_df[calculated_dates_df['date'].dt.year == year]
            
            if year_dates.empty:
                print(f"⚠️  No hay fechas en la BD para el cliente {client_id} en el año {year}")
                return []
            
            # Organizar fechas por actividad
            dates_data = []
            oc_dates = year_dates[year_dates['activity_name'] == 'Fecha Envío OC'].sort_values('date_position')
            entrega_dates = year_dates[year_dates['activity_name'] == 'Fecha Entrega'].sort_values('date_position')
            
            # Combinar fechas de OC y entrega
            max_dates = max(len(oc_dates), len(entrega_dates))
            
            for i in range(max_dates):
                oc_date = None
                entrega_date = None
                
                if i < len(oc_dates):
                    oc_date_obj = oc_dates.iloc[i]['date']
                    oc_date = format_date_spanish(oc_date_obj)
                
                if i < len(entrega_dates):
                    entrega_date_obj = entrega_dates.iloc[i]['date']
                    entrega_date = format_date_spanish(entrega_date_obj)
                
                # Solo añadir si tenemos al menos una fecha
                if oc_date or entrega_date:
                    dates_data.append({
                        'Fecha envío OC': oc_date or '',
                        'Fecha Entrega': entrega_date or ''
                    })
            
            print(f"✅ Obtenidas {len(dates_data)} fechas de la BD para cliente {client_id} año {year}")
            return dates_data
            
        except Exception as e:
            print(f"❌ Error obteniendo fechas de la BD para cliente {client_id}: {e}")
            return []
    
    def create_calendar_table_data(self, client_data, year=2025):
        """
        Crea los datos de la tabla de calendario para el cliente
        """
        # SIEMPRE usar únicamente fechas desde BD - NO usar métodos algorítmicos
        dates = self.generate_dates_from_database(client_data, year)
        
        # Si no hay fechas en BD, retornar vacío (NO generar algorítmicamente)
        if not dates:
            print(f"⚠️  No hay fechas en la BDD para el cliente {client_data.get('id')} en el año {year}")
            dates = []
        
        # Crear estructura de tabla similar a la imagen
        table_data = []
        
        # Cabecera del cliente
        client_name = client_data.get('nombre_cliente', 'Cliente')
        table_data.append([client_name, '', '', ''])
        
        # Cabeceras de columnas si hay fechas
        if dates:
            # Crear pares de columnas para cada semestre
            first_half = []
            second_half = []
            
            for d in dates:
                if d['Fecha envío OC']:  # Solo procesar si hay fecha de OC
                    try:
                        # Convertir fecha para determinar el mes
                        date_str = f"{year}-{d['Fecha envío OC']}"
                        date_obj = datetime.strptime(date_str, '%Y-%d-%b')
                        
                        if date_obj.month <= 6:
                            first_half.append(d)
                        else:
                            second_half.append(d)
                    except (ValueError, TypeError):
                        # Si hay error parseando la fecha, ponerla en primera mitad
                        first_half.append(d)
            
            # Cabeceras
            table_data.append(['Fecha envío OC', 'Fecha Entrega', 'Fecha envío OC', 'Fecha Entrega'])
            
            # Datos - balancear entre primera y segunda mitad del año
            max_rows = max(len(first_half), len(second_half))
            
            for i in range(max_rows):
                row = ['', '', '', '']
                
                # Primera mitad del año
                if i < len(first_half):
                    row[0] = first_half[i]['Fecha envío OC']
                    row[1] = first_half[i]['Fecha Entrega']
                
                # Segunda mitad del año
                if i < len(second_half):
                    row[2] = second_half[i]['Fecha envío OC']
                    row[3] = second_half[i]['Fecha Entrega']
                
                table_data.append(row)
        
        return table_data
    
    def generate_document(self, client_data, output_path=None, year=None):
        """
        Genera un documento Word para un cliente específico
        """
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(f"Plantilla no encontrada: {self.template_path}")
        
        # Cargar plantilla
        doc = Document(self.template_path)
        
        # Obtener nombre del cliente
        client_name = client_data.get('nombre_cliente', 'Cliente')
        safe_client_name = ''.join(c for c in client_name if c.isalnum() or c in (' ', '-', '_'))[:50]
        
        # Reemplazar marcador del nombre del cliente
        for paragraph in doc.paragraphs:
            if 'NOMBRE DEL CLIENTE' in paragraph.text:
                paragraph.text = paragraph.text.replace('NOMBRE DEL CLIENTE', client_name)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Verdana'
        
        # Crear datos de la tabla (pasar año si se especifica)
        if year is not None:
            table_data = self.create_calendar_table_data(client_data, year)
        else:
            table_data = self.create_calendar_table_data(client_data)
        
        # Insertar tabla en el documento
        for paragraph in doc.paragraphs:
            if 'TABLA DE CLIENTE' in paragraph.text:
                # Eliminar el marcador
                p = paragraph._element
                p.getparent().remove(p)
                
                # Crear la tabla
                if table_data:
                    table = doc.add_table(rows=len(table_data), cols=4, style='Table Grid')
                    
                    # Llenar la tabla
                    for i, row_data in enumerate(table_data):
                        for j, cell_value in enumerate(row_data):
                            cell = table.cell(i, j)
                            cell.text = str(cell_value)
                            
                            # Formatear texto
                            for run in cell.paragraphs[0].runs:
                                run.font.name = 'Verdana'
                                run.font.size = Pt(10)
                            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                            # Aplicar colores según el tipo de fila
                            if i == 0:  # Nombre del cliente
                                # Sin color especial para el nombre
                                pass
                            elif i == 1:  # Cabeceras
                                # Color de cabecera (gris claro)
                                shading_elm = OxmlElement('w:shd')
                                shading_elm.set(qn('w:fill'), 'E6E6E6')
                                cell._element.get_or_add_tcPr().append(shading_elm)
                    
                    # Eliminar bordes de la primera fila (excepto borde inferior)
                    for cell in table.rows[0].cells:
                        tcPr = cell._element.get_or_add_tcPr()
                        tcPr.clear()
                        # Borde inferior
                        border = OxmlElement('w:bottom')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '4')
                        border.set(qn('w:space'), '0')
                        tcPr.append(border)
                
                break
        
        # Definir ruta de salida
        if output_path is None:
            output_path = f"CF {safe_client_name} CALENDARIO 2025.docx"
        
        # Guardar documento
        doc.save(output_path)
        return output_path
    
    def generate_multiple_documents(self, clients_data, output_directory=None):
        """
        Genera documentos para múltiples clientes y los comprime en ZIP
        """
        if output_directory is None:
            output_directory = tempfile.mkdtemp()
        
        generated_files = []
        
        for client in clients_data:
            try:
                client_name = client.get('nombre_cliente', 'Cliente')
                safe_name = ''.join(c for c in client_name if c.isalnum() or c in (' ', '-', '_'))[:50]
                output_path = os.path.join(output_directory, f"CF {safe_name} CALENDARIO 2025.docx")
                
                generated_path = self.generate_document(client, output_path)
                generated_files.append(generated_path)
                
            except Exception as e:
                print(f"Error generando calendario para {client.get('nombre_cliente', 'Cliente')}: {e}")
                continue
        
        return generated_files
    
    def create_zip_from_files(self, file_paths, zip_path=None):
        """
        Crea un archivo ZIP con los documentos generados
        """
        if zip_path is None:
            zip_path = os.path.join(tempfile.gettempdir(), f"Calendarios_GL_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip")
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file_path in file_paths:
                if os.path.exists(file_path):
                    # Usar solo el nombre del archivo en el ZIP
                    arcname = os.path.basename(file_path)
                    zipf.write(file_path, arcname)
        
        return zip_path

def get_clients_by_type(tipo_cliente=None, country_filter=None):
    """
    Obtiene clientes filtrados por tipo y/o país
    """
    df_clients = get_clients()
    
    # Convertir DataFrame a lista de diccionarios
    if df_clients.empty:
        return []
    
    clients = df_clients.to_dict('records')
    
    # Renombrar columnas para compatibilidad con la interfaz
    for client in clients:
        # Asegurar que el ID esté disponible
        if 'id' not in client:
            print(f"Advertencia: Cliente {client.get('name', 'Desconocido')} no tiene ID")
        
        # Mapear campos a nombres esperados por la interfaz
        if 'name' in client and 'nombre_cliente' not in client:
            client['nombre_cliente'] = client['name']
        
        # Los campos tipo_cliente y pais ya existen con los nombres correctos
        # Solo añadir si no existen
        if 'tipo_cliente' not in client and 'type' in client:
            client['tipo_cliente'] = client['type']
        if 'pais' not in client and 'country' in client:
            client['pais'] = client['country']
        
        # Inferir frecuencia del calendario_sap si no existe (solo como fallback)
        if 'frecuencia' not in client:
            calendario_sap = client.get('calendario_sap', '')
            if calendario_sap:
                if calendario_sap.startswith('Q') or '15' in calendario_sap:
                    client['frecuencia'] = 'Quincenal'
                elif calendario_sap.startswith('M') or 'M' in calendario_sap:
                    client['frecuencia'] = 'Mensual'
                elif calendario_sap.startswith('B') or '60' in calendario_sap:
                    client['frecuencia'] = 'Bimensual'
                elif calendario_sap.startswith('T') or '90' in calendario_sap:
                    client['frecuencia'] = 'Trimestral'
                elif calendario_sap.startswith('S') or '180' in calendario_sap:
                    client['frecuencia'] = 'Semestral'
                elif calendario_sap.startswith('A') or '365' in calendario_sap:
                    client['frecuencia'] = 'Anual'
                else:
                    client['frecuencia'] = 'Mensual'
            else:
                client['frecuencia'] = 'Mensual'
    
    # Aplicar filtros adicionales si no están ya aplicados por get_clients()
    if tipo_cliente and tipo_cliente != 'Todos':
        clients = [c for c in clients if c.get('tipo_cliente') == tipo_cliente]
    
    if country_filter:
        clients = [c for c in clients if c.get('pais') == country_filter]
    
    return clients

def get_available_client_types():
    """
    Obtiene los tipos de cliente disponibles en la base de datos
    """
    df_clients = get_clients()
    
    if df_clients.empty:
        return []
    
    clients = df_clients.to_dict('records')
    types = set()
    
    for client in clients:
        # Buscar el campo tipo_cliente en diferentes posibles nombres de columna
        client_type = client.get('tipo_cliente') or client.get('type') or client.get('client_type')
        if client_type:
            types.add(client_type)
    
    return sorted(list(types))